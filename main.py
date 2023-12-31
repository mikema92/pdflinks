from files import get_filename, find_files_with_ext
from drive import upload_with_conversion, get_drive_service, create_folder
import argparse
import tkinter as tk
from tkinter import filedialog
import docx

# inputs: case name (drive folder name), folder name (local folder name with pdf files),  
# get all pdf files of a folder
# upload pdf files and save their titles and links and order
# write in word doc (title without order hyperlinked to google docs)

PDF_EXT = 'pdf'
DRIVE_URL_PREFIX = 'https://docs.google.com/file/d/'

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def generateTitlesFromPDFs(case, folder, output):
    files = find_files_with_ext(folder, PDF_EXT)
    if len(files) == 0:
        print(f'no pdf files found in folder {folder}, did you select the correct folder? the pdf files should be in the file you select...')
        return
    service = get_drive_service()
    fid = create_folder(case, service)
    document = docx.Document()
    table = document.add_table(rows=0, cols=2)

    for f in files:
        name = get_filename(f, PDF_EXT)
        id = upload_with_conversion(f, fid, PDF_EXT, service)
        # write in word doc: hyperlinked title without order
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]
        add_hyperlink(p, f"{DRIVE_URL_PREFIX}{id}", name)
    
    document.save(f"{folder}/{output}.docx")
    print(f'document created in {folder}/{output}.docx')

if __name__ == "__main__":
    parser = argparse.ArgumentParser(prog='PDFHyperlinks', description='uploads all pdfs to google drive and produce a word doc with links to the pdfs')
    parser.add_argument('-c', '--case', type=str, help='An optional integer argument')
    parser.add_argument('-f', '--filepath', type=str, help='An optional integer argument')
    parser.add_argument('-o', '--out_filename', type=str, help='An optional integer argument')
    args = parser.parse_args()
    # This code will only be executed 
    # if the script is run as the main program
    file_path = args.filepath 
    case = args.case
    outfname = args.out_filename

    if case == None:
        case = input("Enter new case folder name (only alphabets):\n")
    if outfname == None:
        outfname = input("Enter output name of word document:\n")
    if file_path == None:
        print("Select folder where the pdfs reside, this will also be where the word document is created\n")
        root = tk.Tk()
        root.withdraw()
        file_path = filedialog.askdirectory()
    generateTitlesFromPDFs(case, file_path, outfname)
